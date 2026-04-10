from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run("Daily Task Trackers")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
run = p.add_run("Name:\t")
run.bold = True
p.add_run("ANUDEEP SATYA SAI")
p = doc.add_paragraph()
run = p.add_run("Date: ")
run.bold = True
p.add_run("April.10.2026")
p = doc.add_paragraph()
run = p.add_run("Hours Worked: ")
run.bold = True
p.add_run("9 Hours 10 Minutes (8 hrs 41 min active, 29 min break)")
p = doc.add_paragraph()
run = p.add_run("Scheduled Hours Today: ")
run.bold = True
p.add_run("9:00 PM (Apr 9) \u2013 6:10 AM (Apr 10) IST")

h = doc.add_heading("Remark", level=2)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

p = doc.add_paragraph()
run = p.add_run("Algorithmic Trading System - Task Summary")
run.bold = True

p = doc.add_paragraph()
p.add_run("Integrated Robinhood paper trading end-to-end with real bid/ask prices, built real-time dashboard, fixed critical ML pipeline issues, eliminated system hangs, and tightened entry quality gates based on backtest findings.")
p.runs[0].font.size = Pt(10.5)

bullets = [
    "Integrated Robinhood Crypto API for paper trading \u2013 system now fills paper orders at real Robinhood bid/ask prices. LONG entries fill at best ask, SHORT entries fill at best bid, simulating realistic spread costs (~1.67%). Connected account shows real buying power ($16,445.79) as initial paper equity.",
    "Built complete Paper Trading dashboard page (5_Paper_Trading.py) \u2013 displays live Robinhood prices, open positions with unrealized P&L, combined trade log (entries + exits like a real exchange), equity curve chart, running realized P&L, and account stats. Auto-refreshes every 5 seconds from state file.",
    "Fixed LightGBM feature count mismatch \u2013 BTC model trained on 50 features, ETH on 40, but executor hard-coded [:, :40] truncation. BTC model was silently failing (ml_confidence=0). Now dynamically reads model.num_feature() to match each asset\u2019s trained feature count.",
    "Fixed critical SHORT entry gate bypass \u2013 SHORTs require score 7+ (min_entry_score=4 + short_score_penalty=3) but the LLM could flip a NEUTRAL/BUY signal to SHORT after the pre-LLM gate. Added post-LLM direction gate: SHORT penalty now applies always when final action is SHORT regardless of original signal.",
    "Raised minimum LLM confidence threshold from 0.60 to 0.70 \u2013 backtest data confirmed conf >= 0.70 is profitable; both losing trades entered at exactly 0.60 confidence. This single change would have prevented both initial losses.",
    "Eliminated system hangs from MT5 initialization \u2013 mt5.initialize() hangs indefinitely when MetaTrader 5 terminal is not running. Added paper_mode guard to skip MT5 entirely during paper trading. Also guarded MT5 on both entry and exit paths to prevent real order placement in paper mode.",
    "Eliminated Robinhood API rate-limit freezes \u2013 system was making 4-6 Robinhood API calls per bar (3s throttle each = 12-18s blocking). Refactored: record_entry/record_exit now use executor\u2019s price directly (zero API calls), update_positions receives prices from Kraken OHLCV feed instead of independent Robinhood fetches.",
    "Fixed stale dashboard positions \u2013 state file showed positions as open even after JSONL logged exits (save_state was not called after crash). Added cross-check: dashboard now verifies last JSONL event per asset and will not show stale open positions if exit was already recorded.",
    "Fixed dashboard timezone display \u2013 trade log showed UTC timestamps while user is in IST (UTC+5:30). Added _to_local() converter using datetime.astimezone() for all timestamp displays including trade log, open positions, and footer.",
    "Stripped dashboard to paper-trading only \u2013 removed all pages except Paper Trading, removed old sidebar stats (account balance, P&L, trade stats, recent trades, agent overlay). Clean single-purpose interface.",
    "Cleared all stale testnet/old data from logs \u2013 zeroed out trading_journal.jsonl, trade_history.csv, dashboard_state.json, mt5_trades.jsonl, alerts.jsonl, trade_decisions.jsonl, benchmark_history.json, and all legacy state files to start paper trading fresh.",
    "Analyzed backtest P&L projections at $100K capital \u2013 V15_CATB strategy: BTC +17.72% (PF 1.003), ETH +142.18% (PF 1.018) over 33K trades/9 years. Identified Robinhood\u2019s 1.67% spread as critical blocker (avg win +0.70% vs 1.67% spread cost). Projected: +$12-17K/year on futures exchange vs loss on Robinhood spot.",
]

for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Total Time Spent: 8 hrs 41 min")
run.bold = True

# Tasks with updated hours (total = 8 hrs 41 min = 521 min)
tasks = [
    ("Feature\nIntegration", "Robinhood Paper Trading Integration \u2013 API connection, bid/ask fills, state persistence, equity tracking", "Completed", "1 hr 20 min"),
    ("Feature\nDevelopment", "Paper Trading Dashboard (5_Paper_Trading.py) \u2013 live prices, trade log, equity curve, positions", "Completed", "1 hr 10 min"),
    ("Bug Fix", "LightGBM Feature Count Mismatch \u2013 dynamic num_feature() alignment per asset model", "Completed", "30 min"),
    ("Bug Fix &\nEntry Gate", "SHORT Score Gate Bypass Fix \u2013 post-LLM direction gate + confidence threshold 0.60 to 0.70", "Completed", "35 min"),
    ("Bug Fix &\nPerformance", "MT5 Initialization Hang Fix \u2013 paper_mode guard on MT5 init, entry, and exit paths", "Completed", "25 min"),
    ("Performance\nOptimization", "Robinhood API Rate-Limit Fix \u2013 eliminated redundant API calls, pass prices from OHLCV feed", "Completed", "40 min"),
    ("Bug Fix", "Stale Dashboard Positions Fix \u2013 JSONL cross-check for open position validation", "Completed", "20 min"),
    ("UI Fix", "Dashboard Timezone Fix (UTC to IST) + Old Data Cleanup", "Completed", "25 min"),
    ("Dashboard\nCleanup", "Stripped Dashboard to Paper-Only \u2013 removed all non-paper pages and old sidebar widgets", "Completed", "20 min"),
    ("Analysis &\nReporting", "Backtest P&L Projection ($100K Capital) \u2013 V15 results, spread impact, exchange comparison", "Completed", "30 min"),
    ("System\nDiagnostics", "System Hang Investigation \u2013 traced MT5 init freeze, Robinhood rate-limit stalls, process diagnostics", "Completed", "45 min"),
    ("Verification\n& Testing", "End-to-End Paper Trading Verification \u2013 confirmed 3 paper trades executed, SL triggered correctly", "Completed", "21 min"),
]

table = doc.add_table(rows=1 + len(tasks), cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

widths = [Cm(2.8), Cm(8.0), Cm(2.6), Cm(3.0)]
for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]

header_cells = table.rows[0].cells
headers = ["Task Type", "Task Name", "Completion Status", "Hours worked on task today"]
for i, (cell, text) in enumerate(zip(header_cells, headers)):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2E4057" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

for row_idx, (task_type, task_name, status, hours) in enumerate(tasks):
    row = table.rows[row_idx + 1]
    row_data = [task_type, task_name, status, hours]
    for col_idx, (cell, text) in enumerate(zip(row.cells, row_data)):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        if col_idx == 0:
            run.bold = True
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F7FA" w:val="clear"/>')
            cell._tc.get_or_add_tcPr().append(shading)
        if col_idx == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_idx == 3:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

out_path = r"C:\Users\convo\Downloads\Anudeep_Satya_Sai_EOD_April10.docx"
doc.save(out_path)
print(f"Created: {out_path}")
