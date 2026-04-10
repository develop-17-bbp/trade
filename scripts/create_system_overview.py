"""
Crypto Trading Agent - System Architecture Overview Document
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# Helper functions
def add_title(text, size=24):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def add_subtitle(text, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x77)
    return p

def add_heading(text, level=1):
    return doc.add_heading(text, level=level)

def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        # Shade header
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '1A1A2E'
        })
        shading.append(shd)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()  # spacing
    return table

# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_title('Crypto Trading Agent', 28)
add_title('System Architecture Overview', 20)
doc.add_paragraph()
add_subtitle('EMA(8) Crossover + Multi-LLM + 13-Agent Orchestrator')
add_subtitle('Complete Pipeline Documentation & Future Strategy Integration Guide')
doc.add_paragraph()
add_subtitle('Version 13 (BTC) / Version 14 (ETH)')
add_subtitle('April 9, 2026')
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Components: 85+ Active | 7 ML Models | 13 Agents | 2-Pass LLM | 17+ Indicators')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x77, 0x77, 0x99)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════
add_heading('Table of Contents', 1)
toc_items = [
    '1. System Overview & Architecture',
    '2. Complete Pipeline Flow (10 Stages)',
    '3. Layer-by-Layer Breakdown',
    '   3.1 Data Ingestion Layer',
    '   3.2 Signal Generation Layer',
    '   3.3 Entry Scoring Engine',
    '   3.4 Machine Learning Layer (7 Models)',
    '   3.5 Agent Orchestrator (13 Agents + Debate)',
    '   3.6 LLM Decision Layer (2-Pass Brain)',
    '   3.7 Risk & Protection Layer',
    '   3.8 Order Execution Layer',
    '   3.9 Position Management Layer',
    '   3.10 Monitoring & Feedback Layer',
    '4. Pipeline Health Audit Results',
    '5. Features NOT Yet Implemented (Gaps)',
    '6. Current Strategy: EMA(8) Crossover Deep Dive',
    '7. Future: Multi-Strategy Integration Architecture',
    '8. Vision: LLM as Strategy Neural Network',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.add_run(item).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ════════════════════════════════════════════════════════════════
add_heading('1. System Overview & Architecture', 1)

add_para('This is an algorithmic cryptocurrency trading system that combines a traditional EMA(8) crossover strategy with deep ML enrichment, a 13-agent debate engine, and multi-LLM confirmation before executing any trade. The system runs 24/7 on Bybit and Delta Exchange testnets, with MetaTrader 5 mirroring.')

add_para('Core Philosophy:', bold=True)
add_bullet('', bold_prefix='Signal First: ')
add_para('The EMA(8) crossover generates a BUY/SELL candidate. No crossover = no trade evaluation.')
add_bullet('', bold_prefix='Score Gate: ')
add_para('A 10-point scoring system filters signal quality (min=4, max=7). Bad signals never reach the LLM.')
add_bullet('', bold_prefix='ML Enrichment: ')
add_para('7 ML models add regime context, direction prediction, and quality scores.')
add_bullet('', bold_prefix='Agent Debate: ')
add_para('13 specialized agents vote independently, then debate. Bayesian consensus feeds into LLM.')
add_bullet('', bold_prefix='LLM Confirmation: ')
add_para('2-pass LLM (Mistral scanner + Llama analyst) makes the final GO/NO-GO decision.')
add_bullet('', bold_prefix='Bear Veto: ')
add_para('A contrarian LLM argues AGAINST the trade. If bear_risk >= 9/10, trade is vetoed.')

add_para('Technology Stack:', bold=True)
add_table(
    ['Component', 'Technology', 'Purpose'],
    [
        ['Language', 'Python 3.14', 'Core system'],
        ['Exchange APIs', 'CCXT (Bybit, Delta)', 'Market data + order execution'],
        ['LLM (Scanner)', 'Mistral 7B via Ollama', 'Pattern recognition, 1st pass'],
        ['LLM (Analyst)', 'Llama 3.2 3B via Ollama', 'Risk analysis, 2nd pass'],
        ['LLM (Fallback)', 'Claude Haiku 4.5', 'Cloud fallback when Ollama unavailable'],
        ['ML Framework', 'PyTorch + LightGBM', 'LSTM, PatchTST, binary classifier'],
        ['Broker Mirror', 'MetaTrader 5', 'Visual chart + execution mirroring'],
        ['GPU Server', 'Linux + Cloudflare Tunnel', 'Remote Ollama inference + model training'],
        ['Dashboard', 'Streamlit + FastAPI', 'Real-time monitoring'],
        ['Data Sources', 'Binance, Bybit, LCW, On-Chain', 'OHLCV, sentiment, whale flows'],
    ],
    col_widths=[1.5, 2.5, 3.0]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 2. COMPLETE PIPELINE FLOW
# ════════════════════════════════════════════════════════════════
add_heading('2. Complete Pipeline Flow (10 Stages)', 1)

add_para('Every 10 seconds, for each asset (BTC, ETH), the following pipeline executes:')

stages = [
    ('STAGE 1: Data Fetch', 'executor.py:1550-1561',
     'Fetch 5-minute OHLCV from exchange (100 candles). Also fetch 1m, 15m, 1h, 4h for multi-timeframe analysis. Fetch order book for microstructure (bid/ask imbalance, support/resistance levels).'),

    ('STAGE 2: EMA(8) Signal Detection', 'executor.py:1828-1996',
     'Compute EMA(8) on each timeframe. Detect inflection points (EMA direction change). Identify NEW LINE entries: prior trend 3+ bars, direction flipped, entry within 1-5 bars of inflection. Output: BUY, SELL, or NEUTRAL.'),

    ('STAGE 3: Entry Score Computation', 'executor.py:2087-2232',
     'Score the signal quality on a 0-25 point scale using 10 indicator confirmations: EMA slope strength (0-3), consecutive direction (0-3), price-EMA separation (0-2), candle momentum (0-2), RSI (0-2), ADX (0-2), MACD (0-2), choppiness+MFI (0-2), OBV volume (0-1), S/R levels for ETH (0-2). Gate: score must be 4-7. Below 4 = weak signal, above 7 = momentum trap.'),

    ('STAGE 4: ML Model Enrichment', 'executor.py:2547-2750',
     '7 models add context to ml_context dict: (1) HMM 4-state regime detector - HARD BLOCKS on CRISIS. (2) Kalman trend filter + SNR. (3) GARCH volatility forecast. (4) Hurst exponent - HARD BLOCKS on mean-reverting regime. (5) LightGBM SKIP/TRADE classifier. (6) LSTM Ensemble 3-model direction prediction. (7) RL Agent Q-learning position sizing + SL multiplier.'),

    ('STAGE 5: Agent Orchestrator + Debate', 'executor.py:2891-2920',
     '13 specialized agents run in parallel (ThreadPoolExecutor). Each votes: LONG, SHORT, or FLAT with confidence. Then DebateEngine runs: agents challenge opposing views. AgentCombiner produces Bayesian weighted consensus. Result feeds into LLM prompt as "expert panel" context.'),

    ('STAGE 6: LLM Brain Decision (2-Pass)', 'executor.py:3203-3256',
     'Pass 1 (Mistral 7B): Pattern Scanner - identifies chart formations, momentum patterns, confluence of signals. Pass 2 (Llama 3.2): Risk Analyst - evaluates downside, regime risk, historical pattern outcomes. Combined output: proceed (bool), confidence (0-1), position_size_pct, risk_score, predicted_l_level.'),

    ('STAGE 7: Bear Veto Agent', 'executor.py:3284-3323',
     'Separate LLM call with contrarian prompt: "Argue why this trade will FAIL." Scores bear_risk 1-10. If bear_risk >= 9: HARD VETO (trade cancelled). If bear_risk >= 7: REDUCE position size to 70%. Prevents overconfidence bias in bull-leaning system.'),

    ('STAGE 8: Quality Gates', 'executor.py:3339-3386',
     'Final checks: LLM confidence >= 0.60 threshold. Position size capped at 5% of equity. Exchange-specific minimum order size check. Protections system: drawdown guard, pair lock, ROI table.'),

    ('STAGE 9: Order Execution', 'executor.py:3511-3550',
     'Smart limit order at best bid/ask. If limit rejected: market order fallback. If market rejected: limit with price cap. Fill price verification + slippage monitoring. MT5 mirror/execute for visual tracking.'),

    ('STAGE 10: Position Management', 'executor.py:3720-4328',
     'Trailing stop-loss ratchet: L1 (initial) -> L2 -> L3 -> ... -> L38+. Each level locks in more profit. SL moves to breakeven at L2, then progressively higher. Exit triggers: SL hit, time limit (6h max hold), opposite EMA signal, or manual override.'),
]

for stage_name, file_ref, desc in stages:
    add_heading(stage_name, 3)
    p = doc.add_paragraph()
    r = p.add_run(f'File: {file_ref}')
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x99)
    add_para(desc)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 3. LAYER-BY-LAYER BREAKDOWN
# ════════════════════════════════════════════════════════════════
add_heading('3. Layer-by-Layer Breakdown', 1)

# 3.1 Data
add_heading('3.1 Data Ingestion Layer', 2)
add_table(
    ['Module', 'Source', 'Data Type', 'Status'],
    [
        ['fetcher.py', 'Bybit/Delta CCXT', 'OHLCV, Order Book, Trades', 'ACTIVE'],
        ['free_tier_fetchers.py', 'CoinGecko/Binance/Yahoo', 'Fallback OHLCV', 'ACTIVE'],
        ['on_chain_fetcher.py', 'Blockchain.info/DefiLlama', 'Hashrate, TVL, Whale flows', 'ACTIVE'],
        ['news_fetcher.py', 'NewsAPI', 'Crypto news headlines', 'ACTIVE'],
        ['polymarket_fetcher.py', 'Polymarket API', 'Prediction market odds', 'ACTIVE'],
        ['microstructure.py', 'Exchange order book', 'Bid/ask imbalance, depth', 'ACTIVE'],
        ['tick_ingestor.py', 'Exchange websocket', 'Tick-by-tick for VPIN', 'ACTIVE'],
    ],
    col_widths=[1.8, 1.8, 2.0, 0.8]
)

# 3.2 Signal
add_heading('3.2 Signal Generation Layer', 2)
add_para('The signal generator detects EMA(8) inflection points across multiple timeframes:')
add_table(
    ['Timeframe', 'Purpose', 'Weight'],
    [
        ['5-minute', 'Primary signal (anchor timeframe)', 'Highest'],
        ['1-minute', 'Entry timing refinement', 'Low'],
        ['15-minute', 'Noise filter', 'Medium'],
        ['1-hour', 'Higher timeframe trend', 'High'],
        ['4-hour', 'Major trend direction', 'High'],
    ],
    col_widths=[1.5, 3.5, 1.5]
)
add_para('Signal Types:', bold=True)
add_bullet('', bold_prefix='BUY (New Long): ')
add_para('EMA was falling 3+ bars, then starts rising. Price above EMA. Fresh entry (1-5 bars).')
add_bullet('', bold_prefix='SELL (New Short): ')
add_para('EMA was rising 3+ bars, then starts falling. Price below EMA. Fresh entry (1-5 bars).')
add_bullet('', bold_prefix='NEUTRAL: ')
add_para('No inflection detected. No trade evaluation occurs.')

# 3.3 Scoring
add_heading('3.3 Entry Scoring Engine', 2)
add_para('Each signal is scored 0-25 points. Only scores 4-7 pass to the LLM:')
add_table(
    ['Component', 'Points', 'What It Measures'],
    [
        ['EMA Slope Strength', '0-3', 'Steepness of EMA direction change'],
        ['Consecutive EMA Direction', '0-3', 'How many bars EMA has moved in signal direction'],
        ['Price vs EMA Separation', '0-2', 'Distance from price to EMA (too far = extended)'],
        ['Candle Momentum', '0-2', 'Recent candle body sizes (strong vs weak)'],
        ['RSI Confirmation', '0-2', 'RSI aligns with signal direction'],
        ['ADX Trend Strength', '0-2', 'ADX > 20 confirms trend exists'],
        ['MACD Alignment', '0-2', 'MACD histogram supports signal'],
        ['Choppiness + MFI', '0-2', 'Low choppiness + money flow confirms'],
        ['OBV Volume', '0-1', 'On-balance volume supports direction'],
        ['S/R Levels (ETH only)', '0-2', 'Support/resistance proximity adjustment'],
    ],
    col_widths=[2.2, 0.8, 3.5]
)
add_para('v13 Gate Logic:', bold=True)
add_bullet('Score < 4: LOW SCORE BLOCK (weak signal, skip LLM call)')
add_bullet('Score 4-7: PASS to LLM for evaluation')
add_bullet('Score > 7: HIGH SCORE BLOCK (momentum trap, likely to reverse)')
add_bullet('SHORT signals: +3 penalty (effective minimum = 7 for shorts)')

# 3.4 ML
add_heading('3.4 Machine Learning Layer (7 Models)', 2)
add_table(
    ['Model', 'Type', 'Output', 'Role', 'Hard Block?'],
    [
        ['HMM Regime', '4-state Markov', 'BULL/BEAR/SIDEWAYS/CRISIS', 'Regime detection', 'YES (CRISIS)'],
        ['Kalman Filter', 'Adaptive filter', 'Trend slope + SNR', 'Noise reduction', 'No'],
        ['GARCH(1,1)', 'Volatility', 'Vol forecast + regime', 'Risk sizing', 'No'],
        ['Hurst Exponent', 'Fractal analysis', 'H value (0-1)', 'Trend vs mean-revert', 'YES (H<0.45)'],
        ['LightGBM', 'Binary classifier', 'SKIP/TRADE + confidence', 'Quality pre-filter', 'Advisory'],
        ['LSTM Ensemble', '3-model neural net', 'Direction + confidence', 'Prediction', 'Advisory'],
        ['RL Agent', 'Q-learning', 'Size mult, SL mult', 'Position optimization', 'Sizing only'],
    ],
    col_widths=[1.3, 1.2, 1.8, 1.4, 1.0]
)
add_para('Additionally, PatchTST Transformer is loaded for directional forecasting (UP/DOWN/NEUTRAL + liquidity shock probability).')

# 3.5 Agents
add_heading('3.5 Agent Orchestrator (13 Agents + Debate)', 2)
add_para('The orchestrator runs 13 specialized agents in parallel, then conducts an adversarial debate:')
add_table(
    ['Agent', 'Role', 'Vote'],
    [
        ['Data Integrity Validator', 'Pre-gate: sanitizes quant data', 'PASS/FAIL'],
        ['Market Structure Agent', 'HH/HL/LH/LL + BOS/CHoCH patterns', 'LONG/SHORT/FLAT'],
        ['Regime Intelligence', 'Market regime classification', 'LONG/SHORT/FLAT'],
        ['Mean Reversion Agent', 'Detects mean-reversion setups', 'LONG/SHORT/FLAT'],
        ['Trend Momentum Agent', 'Trend + momentum scoring', 'LONG/SHORT/FLAT'],
        ['Risk Guardian', 'Risk assessment per signal', 'LONG/SHORT/FLAT'],
        ['Sentiment Decoder', 'News + social sentiment analysis', 'LONG/SHORT/FLAT'],
        ['Trade Timing Agent', 'Optimal entry timing', 'LONG/SHORT/FLAT'],
        ['Portfolio Optimizer', 'Cross-asset position sizing', 'LONG/SHORT/FLAT'],
        ['Pattern Matcher', 'Historical pattern matching', 'LONG/SHORT/FLAT'],
        ['Loss Prevention Guardian', 'Loss mitigation strategies', 'LONG/SHORT/FLAT'],
        ['Polymarket Agent', 'Prediction market odds', 'LONG/SHORT/FLAT'],
        ['Decision Auditor', 'Post-gate: contradiction check', 'APPROVE/REJECT'],
    ],
    col_widths=[2.0, 2.8, 1.5]
)
add_para('Debate Engine Flow:', bold=True)
add_bullet('Step 1: All 11 analysis agents vote independently (parallel ThreadPoolExecutor)')
add_bullet('Step 2: DebateEngine - agents challenge opposing views, can flip positions')
add_bullet('Step 3: AgentCombiner - Bayesian weighted consensus (weighted by historical accuracy)')
add_bullet('Step 4: Result injected into LLM prompt as "expert panel consensus"')

# 3.6 LLM
add_heading('3.6 LLM Decision Layer (2-Pass Brain)', 2)
add_para('The Trading Brain v2.1 uses a 2-pass architecture:')
add_para('Pass 1 - Pattern Scanner (Mistral 7B):', bold=True)
add_bullet('Receives: candle data, EMA values, entry score, indicator context, agent consensus')
add_bullet('Task: Identify chart formations, momentum patterns, confluence signals')
add_bullet('Output: Pattern analysis + initial confidence')

add_para('Pass 2 - Risk Analyst (Llama 3.2 3B):', bold=True)
add_bullet('Receives: Pass 1 output + ML context (HMM regime, GARCH vol, Hurst, LSTM pred)')
add_bullet('Task: Evaluate downside risk, regime suitability, historical pattern outcomes')
add_bullet('Output: Final decision (proceed/skip), confidence, position size, predicted L-level')

add_para('Additional LLM Features:', bold=True)
add_bullet('Memory Vault: Semantic search of past winning trades for pattern matching')
add_bullet('Winner DNA: Statistical profile of L4+ winning trades (avg confidence, duration)')
add_bullet('Confidence Calibration: Adjusts LLM confidence based on historical accuracy')
add_bullet('Kelly Criterion: Optimal position sizing from calibrated edge estimate')

# 3.7 Risk
add_heading('3.7 Risk & Protection Layer', 2)
add_table(
    ['Protection', 'Type', 'Threshold', 'Action'],
    [
        ['Bear Veto Agent', 'Contrarian LLM', 'bear_risk >= 9', 'HARD VETO'],
        ['Bear Reduce', 'Contrarian LLM', 'bear_risk >= 7', 'Reduce to 70%'],
        ['Confidence Gate', 'LLM output', 'conf < 0.60', 'REJECT'],
        ['Max Drawdown', 'Portfolio', 'daily_loss > 2%', 'HALT trading'],
        ['Pair Lock', 'Per-asset', '3 consecutive losses', 'LOCK 30min'],
        ['ROI Table', 'Time-based', 'Expected ROI vs actual', 'CLOSE if below'],
        ['SL Guard', 'Per-trade', 'SL moved against position', 'BLOCK modification'],
        ['VPIN Guard', 'Flow toxicity', 'VPIN > 0.7', 'WARNING'],
        ['Position Cap', 'Per-trade', '5% of equity', 'CAP size'],
        ['ATR Volatility', 'Per-entry', 'ATR > 2x average', 'BLOCK entry'],
    ],
    col_widths=[1.5, 1.2, 1.8, 1.8]
)

# 3.8 Execution
add_heading('3.8 Order Execution Layer', 2)
add_para('Smart order routing with multiple fallbacks:')
add_bullet('Primary: Limit order at best bid/ask (reduces slippage)')
add_bullet('Fallback 1: Market order if limit rejected')
add_bullet('Fallback 2: Limit with price cap if market rejected')
add_bullet('MT5 Bridge: Mirror/Execute mode for MetaTrader 5 (visual tracking + external broker)')
add_bullet('Multi-Exchange: Bybit + Delta running in parallel threads')

# 3.9 Position
add_heading('3.9 Position Management Layer', 2)
add_para('Trailing Stop-Loss Ratchet System (L-Levels):', bold=True)
add_bullet('L1: Initial stop-loss at entry (based on ATR)')
add_bullet('L2: Move SL to breakeven when trade reaches 1x ATR profit')
add_bullet('L3-L6: Progressive profit lock (SL follows price at increasing intervals)')
add_bullet('L6+: Maximum trailing - locks in significant profit')
add_bullet('L38+: Rare ultra-runner - system lets it ride with very tight trail')
add_para('Exit Triggers:', bold=True)
add_bullet('SL hit at any L-level')
add_bullet('Time limit: 6 hours max hold (configurable)')
add_bullet('Opposite EMA signal (new inflection in opposite direction)')
add_bullet('Manual override via dashboard')

# 3.10 Monitoring
add_heading('3.10 Monitoring & Feedback Layer', 2)
add_table(
    ['Component', 'Purpose', 'Status'],
    [
        ['Trade Journal (JSONL)', 'Every trade logged with full reasoning', 'ACTIVE'],
        ['Health Checker', 'CPU, memory, API responsiveness', 'ACTIVE'],
        ['Drift Detector', 'Model performance vs baseline', 'ACTIVE'],
        ['Auto Healer', 'Self-fixing for common errors', 'ACTIVE'],
        ['Event Guard', 'Flash crash / unusual volume detection', 'ACTIVE'],
        ['Daily Report Generator', 'EOD performance summaries', 'ACTIVE'],
        ['Streamlit Dashboard', '4-page real-time monitoring', 'ACTIVE'],
        ['FastAPI Server', 'REST API for system state', 'ACTIVE'],
    ],
    col_widths=[2.0, 3.0, 1.0]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 4. PIPELINE HEALTH AUDIT
# ════════════════════════════════════════════════════════════════
add_heading('4. Pipeline Health Audit Results', 1)
add_para('Full codebase audit completed April 9, 2026. Pipeline integrity: A- (94/100)', bold=True)

add_heading('Fully Connected & Working', 2)
add_table(
    ['Pipeline Stage', 'Status', 'Grade'],
    [
        ['Data Layer (OHLCV, orderbook, on-chain)', 'All sources wired', 'A'],
        ['Signal Generation (EMA inflection)', 'Complete', 'A'],
        ['Entry Scoring (10-point system)', 'All indicators computed', 'A'],
        ['HMM/Kalman/GARCH/Hurst', 'Hard-integrated, blocks on crisis', 'A'],
        ['LightGBM + LSTM Ensemble', 'Loaded, predictions advisory', 'B+'],
        ['RL Agent (Q-learning)', 'Wired for sizing + SL multiplier', 'A'],
        ['13-Agent Orchestrator + Debate', 'All agents initialized, parallel', 'A'],
        ['Trading Brain v2 (Mistral+Llama)', 'Full 2-pass flow connected', 'A'],
        ['Bear Veto Agent', 'Hard blocks at 9/10', 'A'],
        ['Trade Protections', 'SL guard, drawdown, pair lock active', 'A'],
        ['Order Execution', 'Smart limit + fallbacks', 'A'],
        ['Position Management', 'Trailing SL L1-L38+', 'A'],
        ['Journal Logging', 'Every trade recorded', 'A'],
    ],
    col_widths=[3.0, 2.5, 0.8]
)

add_heading('Disconnected / Dead Code', 2)
add_table(
    ['Component', 'Issue', 'Impact'],
    [
        ['PatchTST per-asset', 'Loaded shared model only, BTC/ETH specific unused', 'Low'],
        ['Alpha Decay Model', 'Imported but predict() never called in executor', 'Low'],
        ['ProfitProtector', 'Redundant with protections.py', 'None'],
        ['Alerting Module', 'Available but not wired into execution', 'Medium'],
        ['Meta Controller', 'Dashboard only, not in main loop', 'Low'],
        ['Adaptive Engine', 'Imported but select_best_strategy() not called', 'Medium'],
    ],
    col_widths=[1.8, 3.0, 1.0]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 5. FEATURES NOT IMPLEMENTED
# ════════════════════════════════════════════════════════════════
add_heading('5. Features NOT Yet Implemented (Gaps)', 1)

add_heading('Critical Gaps', 2)
add_table(
    ['Feature', 'Current State', 'Impact', 'Effort'],
    [
        ['Auto-Retraining Pipeline', 'Manual via scripts only', 'Models go stale over time', 'Medium'],
        ['Multi-Strategy Switching', 'EMA(8) hardcoded, AdaptiveEngine not wired', 'Cannot adapt to regime changes', 'High'],
        ['Bybit Testnet Price Fix', 'Testnet returns fantasy prices ($300K-$1.5M BTC)', 'Journal PnL meaningless', 'Medium'],
        ['LightGBM/LSTM Enforcement', 'Models predict but dont block', 'ML models are advisory only', 'Low'],
        ['Duplicate Trade Fix', 'Same trade logged 2-3 times', 'Journal stats inflated', 'Low'],
        ['Delta STUCK Fix', 'False no-liquidity detection', 'Profitable positions reported as stuck', 'Low'],
    ],
    col_widths=[2.0, 2.5, 2.0, 0.8]
)

add_heading('Nice-to-Have Gaps', 2)
add_table(
    ['Feature', 'Current State', 'Notes'],
    [
        ['NBeats/TFT Forecasters', 'Throw NotImplementedError', 'Placeholder code exists'],
        ['FinGPT Forecaster', 'Requires unavailable SDK', 'Would need custom integration'],
        ['Pairs/Spread Trading', 'Cointegration engine built, not wired', 'BTC-ETH spread strategy'],
        ['Options Hedging', 'Portfolio hedger framework only', 'Would protect against tail risk'],
        ['Premium News (Bloomberg)', 'Only free APIs', 'Better sentiment accuracy'],
        ['More Exchanges', 'Only Bybit + Delta', 'Binance Futures, OKX, Kraken'],
        ['Monte Carlo VaR', 'Code present, not wired', 'Better risk quantification'],
        ['LoRA Fine-tuning', 'Placeholder exists', 'Fine-tune Mistral on winning trades'],
        ['Temporal Transformer', 'Code present, not used', 'Alternative to PatchTST'],
    ],
    col_widths=[2.0, 2.5, 2.5]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 6. CURRENT STRATEGY DEEP DIVE
# ════════════════════════════════════════════════════════════════
add_heading('6. Current Strategy: EMA(8) Crossover Deep Dive', 1)

add_para('The system currently runs a single strategy: EMA(8) Crossover with LLM confirmation.')

add_heading('What is EMA(8) Crossover?', 2)
add_para('An Exponential Moving Average with period 8 is computed on the close prices. When the EMA changes direction (inflection point), a trade signal is generated:')
add_bullet('EMA was falling for 3+ bars, starts rising -> BUY signal')
add_bullet('EMA was rising for 3+ bars, starts falling -> SELL signal')
add_bullet('The "new line" must be fresh (1-5 bars old) to prevent late entries')

add_heading('Why EMA(8)?', 2)
add_bullet('Fast enough to catch 5-minute trend changes')
add_bullet('Slow enough to filter out 1-minute noise')
add_bullet('Inflection detection gives early entries (before price confirms)')
add_bullet('10-year backtest: PF=1.003 BTC, PF=1.019 ETH (with all filters)')

add_heading('Strategy Parameters (v13 BTC / v14 ETH)', 2)
add_table(
    ['Parameter', 'BTC (v13)', 'ETH (v14)', 'Purpose'],
    [
        ['ema_period', '8', '8', 'EMA calculation period'],
        ['min_entry_score', '4', '4', 'Minimum score to evaluate'],
        ['max_entry_score', '7', '7', 'Maximum score (momentum trap cap)'],
        ['short_score_penalty', '3', '3', 'Extra score needed for SHORTs'],
        ['S/R Levels', 'DISABLED', 'ENABLED', 'Support/resistance scoring'],
        ['struct_window', '5', '5', 'Dynamic SL structure lookback'],
    ],
    col_widths=[1.8, 1.0, 1.0, 3.0]
)

add_heading('What Makes It Different From Simple EMA Crossover', 2)
add_para('A simple EMA crossover would have PF < 0.95 (losing money). This system adds 9 layers of intelligence on top:')
add_bullet('', bold_prefix='Layer 1 (Score Gate): ')
add_para('10-point scoring filters out 60%+ of signals before LLM call')
add_bullet('', bold_prefix='Layer 2 (ML Regime): ')
add_para('HMM blocks trades during market crisis. Hurst blocks in choppy markets.')
add_bullet('', bold_prefix='Layer 3 (LSTM/LightGBM): ')
add_para('Neural nets predict if this specific setup will reach L2+ or die at L1')
add_bullet('', bold_prefix='Layer 4 (Agent Debate): ')
add_para('13 agents argue for/against. Bayesian consensus removes group-think.')
add_bullet('', bold_prefix='Layer 5 (LLM Brain): ')
add_para('Mistral + Llama 2-pass analysis with full market context')
add_bullet('', bold_prefix='Layer 6 (Bear Veto): ')
add_para('Contrarian agent argues against every trade. Blocks overconfident entries.')
add_bullet('', bold_prefix='Layer 7 (Protections): ')
add_para('Drawdown limits, pair locks, ROI table prevent systematic blowups')
add_bullet('', bold_prefix='Layer 8 (Smart Execution): ')
add_para('Limit orders, slippage monitoring, MT5 mirroring')
add_bullet('', bold_prefix='Layer 9 (Trailing SL): ')
add_para('L-level ratchet locks profit progressively, never gives back gains')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 7. MULTI-STRATEGY INTEGRATION
# ════════════════════════════════════════════════════════════════
add_heading('7. Future: Multi-Strategy Integration Architecture', 1)

add_para('The system has partial infrastructure for multiple strategies but currently only runs EMA(8). Here is the architecture for adding new strategies:')

add_heading('Existing Plugin Points', 2)
add_table(
    ['File', 'What Exists', 'Status'],
    [
        ['sub_strategies.py', 'SubStrategy base class with generate_signal() interface', 'Framework only'],
        ['adaptive_engine.py', '5 strategy slots registered (EMA, mean_reversion, trend, vol_breakout, scalping)', 'Registered but not called'],
        ['adaptive_engine.py', 'select_best_strategy() method - chooses strategy per market regime', 'Not wired to executor'],
    ],
    col_widths=[2.0, 3.5, 1.5]
)

add_heading('How To Add a New Strategy (e.g., RSI Mean Reversion)', 2)
add_para('Step 1: Create the Strategy Class', bold=True)
add_para('In sub_strategies.py, extend SubStrategy:')
add_para('class RSIMeanReversion(SubStrategy):', italic=True)
add_para('  def generate_signal(self, prices, highs, lows, volumes) -> int:', italic=True)
add_para('    rsi_val = rsi(prices, 14)', italic=True)
add_para('    if rsi_val[-1] < 30: return 1  # BUY (oversold)', italic=True)
add_para('    if rsi_val[-1] > 70: return -1  # SELL (overbought)', italic=True)
add_para('    return 0  # NEUTRAL', italic=True)

add_para('Step 2: Register in AdaptiveEngine', bold=True)
add_para('Add to adaptive_engine.py strategies dict:')
add_para("strategies['rsi_mean_reversion'] = RSIMeanReversion(period=14)", italic=True)

add_para('Step 3: Add Selection Logic', bold=True)
add_para('In select_best_strategy(), add regime-based selection:')
add_para("if hurst_regime == 'mean_reverting': return 'rsi_mean_reversion'", italic=True)
add_para("if hurst_regime == 'trending': return 'ema_crossover'", italic=True)

add_para('Step 4: Wire AdaptiveEngine into Executor', bold=True)
add_para('Replace hardcoded _compute_tf_signal() call with:')
add_para('strategy = self.adaptive_engine.select_best_strategy(regime_context)', italic=True)
add_para('signal = self.adaptive_engine.strategies[strategy].generate_signal(...)', italic=True)

add_para('Step 5: Retrain ML Models', bold=True)
add_para('LightGBM and LSTM use EMA-specific features (35-39). For RSI strategy:')
add_bullet('Replace features 35-39 with RSI inflection features')
add_bullet('Retrain on historical RSI crossover entries + L-level outcomes')
add_bullet('Create separate model files: lgbm_btc_rsi.txt, lstm_ensemble_btc_rsi/')

add_para('Step 6: Update Entry Scoring', bold=True)
add_para('Each strategy needs its own scoring logic. The current 10-point system is EMA-specific.')

add_heading('Coupling Assessment', 2)
add_table(
    ['Component', 'Strategy-Coupled?', 'Reuse for New Strategy?'],
    [
        ['Signal Detection', 'YES (EMA inflection)', 'Must rewrite per strategy'],
        ['Entry Scoring', 'YES (EMA-specific points)', 'Must create new scoring'],
        ['LightGBM/LSTM', 'YES (Features 35-39)', 'Must retrain'],
        ['HMM/Hurst/GARCH', 'NO (regime detection)', 'Fully reusable'],
        ['Agent Orchestrator', 'NO (generic analysis)', 'Fully reusable'],
        ['LLM Brain', 'NO (reads any signal)', 'Fully reusable'],
        ['Bear Veto', 'NO (contrarian)', 'Fully reusable'],
        ['Protections', 'NO (risk management)', 'Fully reusable'],
        ['Position Management', 'NO (SL ratchet)', 'Fully reusable'],
        ['PatchTST', 'PARTIALLY (price direction)', 'Mostly reusable'],
    ],
    col_widths=[2.0, 1.8, 2.8]
)
add_para('Summary: 60% of the system is strategy-agnostic and immediately reusable. The other 40% (signal detection, scoring, ML feature engineering) needs per-strategy implementation.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 8. VISION: LLM AS STRATEGY NEURAL NETWORK
# ════════════════════════════════════════════════════════════════
add_heading('8. Vision: LLM as Strategy Neural Network', 1)

add_para('Current Architecture: One Strategy, LLM Confirms', bold=True, size=12)
add_para('Right now, the LLM is a "confirmation gate" - the EMA(8) strategy generates signals, and the LLM decides whether to take them. The LLM cannot generate its own signals or switch strategies.')

add_para('Proposed Architecture: LLM as Strategy Selector + Generator', bold=True, size=12)
add_para('The idea is to treat the LLM like a neural network that has "learned" multiple strategy patterns. Just as a neural net has layers that recognize different features, the LLM would have strategy templates loaded as context, and it would dynamically select and combine strategies based on real-time market conditions.')

add_heading('How It Would Work', 2)
add_para('Phase 1: Strategy Pattern Library', bold=True)
add_para('Create a structured library of strategy patterns that the LLM can reference:')
add_bullet('EMA Crossover patterns (current) - works best in trending markets')
add_bullet('RSI Mean Reversion patterns - works best in ranging/choppy markets')
add_bullet('Bollinger Band Squeeze - works best before breakouts')
add_bullet('MACD Divergence - works best at trend reversals')
add_bullet('Volume Profile / VWAP - works best for institutional order flow')
add_bullet('Ichimoku Cloud - works best for multi-timeframe trend confirmation')
add_bullet('Order Flow / Footprint - works best for HFT-style scalping')

add_para('Phase 2: Strategy Templates as "Layers"', bold=True)
add_para('Each strategy becomes a "layer" in the LLM prompt - similar to how neural net layers process different features:')
add_bullet('', bold_prefix='Input Layer: ')
add_para('Raw OHLCV + all indicator values + regime context')
add_bullet('', bold_prefix='Strategy Layers: ')
add_para('Each strategy template analyzes the same data from its perspective')
add_bullet('', bold_prefix='Attention Layer: ')
add_para('LLM "attends" to the most relevant strategy for current conditions')
add_bullet('', bold_prefix='Output Layer: ')
add_para('Combined signal with confidence and position sizing')

add_para('Phase 3: Dynamic Strategy Mixing', bold=True)
add_para('Instead of choosing ONE strategy, the LLM would weight multiple strategies:')
add_para('Example in trending market:', italic=True)
add_bullet('EMA Crossover: 60% weight (primary trend signal)')
add_bullet('MACD Divergence: 20% weight (reversal warning)')
add_bullet('Volume Profile: 15% weight (institutional confirmation)')
add_bullet('RSI: 5% weight (overbought/oversold context)')

add_para('Example in ranging market:', italic=True)
add_bullet('RSI Mean Reversion: 50% weight (primary signal)')
add_bullet('Bollinger Squeeze: 25% weight (breakout detection)')
add_bullet('EMA Crossover: 15% weight (trend start detection)')
add_bullet('Order Flow: 10% weight (smart money tracking)')

add_heading('Implementation Architecture', 2)
add_para('The system would evolve from a single fixed strategy to a dynamic multi-strategy engine:')

add_table(
    ['Component', 'Current', 'Future'],
    [
        ['Signal Source', 'EMA(8) only', 'All strategies compute signals simultaneously'],
        ['Strategy Selection', 'Hardcoded EMA', 'LLM selects based on regime + pattern match'],
        ['Entry Scoring', 'EMA-specific 10-point', 'Per-strategy scoring + combined meta-score'],
        ['ML Models', 'Trained on EMA patterns', 'Trained per-strategy + meta-model'],
        ['LLM Role', 'Confirm/reject EMA signal', 'Select strategy + generate signal + confirm'],
        ['Agent Debate', 'Generic analysis', 'Strategy-aware (agents specialize per strategy)'],
        ['Position Mgmt', 'Same SL ratchet for all', 'Strategy-specific exit rules'],
    ],
    col_widths=[1.5, 2.5, 3.0]
)

add_heading('The Neural Net Analogy', 2)
add_para('In a traditional neural network:')
add_bullet('Input neurons receive raw data (prices, volumes)')
add_bullet('Hidden layers extract features (patterns, trends, support/resistance)')
add_bullet('Different neurons specialize in different pattern types')
add_bullet('The network learns which neurons to weight for each situation')
add_bullet('Output neurons produce the final prediction')

add_para('In the LLM-as-strategy-network:')
add_bullet('', bold_prefix='Input = ')
add_para('Raw OHLCV + all indicators (same as neural net inputs)')
add_bullet('', bold_prefix='Strategy Templates = Hidden Layers. ')
add_para('Each strategy is a "neuron" that processes data from its perspective')
add_bullet('', bold_prefix='Regime Context = Attention Mechanism. ')
add_para('HMM/Hurst/GARCH tell the LLM which "neurons" to weight higher')
add_bullet('', bold_prefix='LLM Reasoning = Backpropagation. ')
add_para('The LLM learns from trade outcomes (memory vault, winner DNA) to refine weights')
add_bullet('', bold_prefix='Final Decision = Output Layer. ')
add_para('Combined weighted signal with confidence')

add_heading('Key Advantage Over Traditional Multi-Strategy Systems', 2)
add_para('Traditional systems use rigid regime detection to switch strategies (if trending, use trend strategy; if ranging, use mean reversion). This fails because:')
add_bullet('Regime changes are detected AFTER the fact (lagging)')
add_bullet('Markets can exhibit multiple regimes simultaneously')
add_bullet('Transitions between regimes are the most profitable moments')

add_para('The LLM approach is superior because:')
add_bullet('LLMs can reason about ambiguous situations ("it looks like a trend but volume is declining")')
add_bullet('LLMs can weight multiple strategies simultaneously (not binary switching)')
add_bullet('LLMs can use qualitative context (news, on-chain, prediction markets)')
add_bullet('LLMs improve over time through memory and feedback (winner DNA extraction)')
add_bullet('LLMs can explain their reasoning (audit trail for each trade)')

add_heading('Roadmap', 2)
add_table(
    ['Phase', 'What', 'When', 'Effort'],
    [
        ['Phase 0', 'Current system: EMA(8) + LLM confirmation', 'NOW', 'Done'],
        ['Phase 1', 'Wire AdaptiveEngine: EMA + RSI strategies switchable', 'Next', '2-3 days'],
        ['Phase 2', 'Add 3-4 more strategy templates + per-strategy ML models', 'After Phase 1', '1-2 weeks'],
        ['Phase 3', 'LLM as strategy selector (regime-aware prompt engineering)', 'After Phase 2', '1 week'],
        ['Phase 4', 'LLM as strategy mixer (weighted multi-strategy signals)', 'After Phase 3', '1-2 weeks'],
        ['Phase 5', 'LoRA fine-tuning on winning strategy selections', 'After Phase 4', '1 week'],
        ['Phase 6', 'Full autonomous strategy generation from LLM', 'Long-term', 'Research'],
    ],
    col_widths=[1.0, 3.5, 1.2, 1.0]
)

doc.add_paragraph()
add_para('End of Document', bold=True)
p = doc.add_paragraph()
r = p.add_run('Generated April 9, 2026 | Crypto Trading Agent v13/v14')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0xAA)

# Save
output_path = r'C:\Users\convo\Downloads\Crypto_Trading_Agent_System_Overview.docx'
doc.save(output_path)
print(f'SAVED: {output_path}')
