from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Daily Task Trackers')
run.bold = True
run.font.size = Pt(16)

# Name/Date/Hours
info = doc.add_paragraph()
r = info.add_run('Name:\t')
r.bold = True
info.add_run('ANUDEEP SATYA SAI\n')
r = info.add_run('Date: ')
r.bold = True
info.add_run('April.08.2026\n')
r = info.add_run('Hours Worked: ')
r.bold = True
info.add_run('11 Hours (10 hrs 6 min active, 54 min break)\n')
r = info.add_run('Scheduled Hours Today: ')
r.bold = True
info.add_run('9:30 pm (Apr 7) - 8:30 am (Apr 8) IST')

# Remark
doc.add_paragraph('Remark', style='Heading 2')

# Task Table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Task Type', 'Task Name', 'Completion Status', 'Hours worked on task today']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

tasks = [
    [
        'Bug Fix &\nInvestigation',
        'BTC Profit Factor Regression Investigation - Traced PF drop from 1.003 to 0.99 across 3 backtest versions (v12, v13, v14). Identified 3 root causes: max_entry_score defaulting to 99 instead of 7, short_score_penalty defaulting to 0 instead of 3, and S/R scoring applied to all assets instead of ETH only',
        'Completed',
        '1 hr 30 min'
    ],
    [
        'Bug Fix &\nCode Repair',
        'Backtest Engine Defaults Fix (full_engine.py) - Fixed max_entry_score default from 99 to 7 (momentum trap cap) and short_score_penalty from 0 to 3. Stored config reference and passed asset+config to signal generator for per-asset behavior',
        'Completed',
        '30 min'
    ],
    [
        'Bug Fix &\nCode Repair',
        'Signal Generator S/R Per-Asset Fix (signal_generator.py) - Added asset and config parameters to compute_entry_score(). Wrapped S/R level scoring in sr_assets check so BTC uses v13 (no S/R) and ETH uses v14 (with S/R). Restored BTC PF from 0.989 to 1.003',
        'Completed',
        '40 min'
    ],
    [
        'Configuration\nAlignment',
        'v13/v14 Parameter Alignment Verification - Audited all code paths (config.yaml, full_engine.py, signal_generator.py, executor.py, backtest CLI, optimizer) to ensure min_entry_score=4, max_entry_score=7, short_score_penalty=3, sr_assets=[ETH] are consistent everywhere',
        'Completed',
        '35 min'
    ],
    [
        'Bug Fix &\nModel Repair',
        'LightGBM Model Corruption Fix - Git merge mangled binary LightGBM model files causing "Model format error" on startup. Retrained both BTC (126 trees, 440KB) and ETH (33 trees, 120KB) LightGBM classifiers locally with fresh market data',
        'Completed',
        '50 min'
    ],
    [
        'Bug Fix &\nIntegration',
        'RL Agent Load Error Fix (executor.py) - Fixed "EMAStrategyRL has no attribute load" error. Changed from .load() method call to constructor pattern: EMAStrategyRL({"rl_model_path": path}). Wired RL Agent decide() into ml_context with action, size_mult, sl_mult, quality fields',
        'Completed',
        '35 min'
    ],
    [
        'Bug Fix &\nIntegration',
        'Bybit Symbol Format Fix (executor.py) - Fixed "bybit does not have market symbol BTC" error. Removed MT5 symbol override from _get_symbol() so Bybit uses CCXT format (BTC/USDT:USDT) instead of bare asset name',
        'Completed',
        '25 min'
    ],
    [
        'Performance\nOptimization',
        'GARCH Model Pre-initialization (executor.py) - Changed from ad-hoc pickle load per tick to pre-loaded in __init__. Both BTC and ETH GARCH models stored in self._garch_per_asset dict, eliminating repeated file I/O during live trading',
        'Completed',
        '20 min'
    ],
    [
        'ML Model\nTraining',
        'ETH Model Training (6/6 models) - Trained all ETH ML models: LightGBM (SKIP/TRADE classifier), HMM (4-state regime detector), GARCH(1,1) volatility, Alpha Decay (half-life=41.6 bars), RL Agent (300 Q-table states), LSTM Ensemble (3 models: LSTM+GRU+BiLSTM)',
        'Completed',
        '1 hr 15 min'
    ],
    [
        'DevOps &\nGit Maintenance',
        'Git History Cleanup - Removed Co-Authored-By tags (Claude Opus 4.6, Claude Sonnet 4.6 variants) from all git commits using git-filter-repo with regex callback. Resolved git merge conflicts on binary model files from Linux GPU server sync',
        'Completed',
        '40 min'
    ],
    [
        'Verification\n& Validation',
        'Full Model Health Check - Verified all 14 ML models (7 per asset x 2) load correctly: LightGBM, LSTM Ensemble, PatchTST, HMM, GARCH, Alpha Decay, RL Agent. Confirmed 13-agent orchestrator, bear veto, and all indicators initialize properly',
        'Completed',
        '25 min'
    ],
    [
        'System\nValidation',
        'Live System Restart & Validation - Restarted trading system, confirmed Bybit OHLCV fetching working (no more symbol errors), verified score filtering (BTC score=8 blocked as momentum trap, ETH score=-3 blocked as low score), all models active, bear veto engaged',
        'Completed',
        '20 min'
    ],
    [
        'Full Codebase\nAudit',
        'System-Wide Status Check - Ran 4 parallel audit agents analyzing executor, ML models, AI/LLM layer, and infrastructure. Identified all misalignments, missing wiring, and runtime errors across the entire trading codebase',
        'Completed',
        '45 min'
    ],
    [
        'Meeting &\nDiscussion',
        'Project Review Meeting with Dr. Sajan Sir - Discussed trading system progress, architecture decisions, model performance, and roadmap for upcoming improvements (5:30 AM - 8:30 AM IST)',
        'Completed',
        '3 hr 00 min'
    ],
]

for task in tasks:
    row = table.add_row()
    for i, val in enumerate(task):
        row.cells[i].text = val

for row in table.rows:
    row.cells[0].width = Inches(1.2)
    row.cells[1].width = Inches(3.8)
    row.cells[2].width = Inches(0.9)
    row.cells[3].width = Inches(1.1)

# Summary
doc.add_paragraph()
summary_heading = doc.add_paragraph()
run = summary_heading.add_run('Algorithmic Trading System - Task Summary')
run.bold = True
run.font.size = Pt(12)

summary = doc.add_paragraph()
r = summary.add_run(
    'Investigated and fixed BTC profit factor regression from 1.003 to 0.99, '
    'repaired corrupted ML models after git merge, wired RL Agent into live executor, '
    'and aligned all v13/v14 parameters across the full codebase.'
)
r.italic = True

doc.add_paragraph('Key highlights:', style='List Bullet')

highlights = [
    'Traced BTC PF regression through 3 backtest versions to identify 3 root causes: wrong defaults in backtest engine (max_entry_score=99, short_score_penalty=0) and S/R scoring applied to BTC (should be ETH only).',
    'Fixed backtest engine defaults and signal generator to respect per-asset S/R configuration via sr_assets parameter. Restored BTC PF from 0.989 back to 1.003.',
    'Repaired LightGBM model corruption caused by git text-merging binary model files. Retrained both BTC (126 trees) and ETH (33 trees) classifiers.',
    'Fixed RL Agent integration error (no public load() method) by using constructor pattern. Wired RL decide() output into ml_context for live trading decisions.',
    'Fixed Bybit symbol format error where MT5 bridge was overriding CCXT symbol format. Bybit now correctly uses BTC/USDT:USDT.',
    'Pre-initialized GARCH models in executor __init__ to eliminate per-tick file I/O overhead.',
    'Trained all 6 ETH ML models: LightGBM, HMM, GARCH, Alpha Decay, RL Agent (300 states), LSTM Ensemble (3 sub-models).',
    'Verified all 14 ML models healthy across BTC and ETH. Confirmed v13/v14 parameters consistent in config.yaml, executor.py, backtest engine, and signal generator.',
    'Cleaned git history to remove Co-Authored-By tags and resolved binary merge conflicts from GPU server sync.',
    'Validated live system restart: Bybit OHLCV working, correct score filtering active (momentum trap cap at 7, short penalty of 3), all 13 agents + bear veto engaged.',
]

for h in highlights:
    doc.add_paragraph(h, style='List Bullet')

total = doc.add_paragraph()
r = total.add_run('Total Time Spent: 10 hr 6 min')
r.bold = True

doc.save(r'C:\Users\convo\Downloads\Anudeep_Satya_Sai_EOD_April8.docx')
print('SAVED: C:\\Users\\convo\\Downloads\\Anudeep_Satya_Sai_EOD_April8.docx')
