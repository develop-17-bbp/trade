"""Generate EOD Daily Tracker for April 9, 2026 session."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Style defaults ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Header ──
p = doc.add_paragraph()
run = p.add_run('Daily Task Trackers')
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
run = p.add_run('Name:\tANUDEEP SATYA SAI\n')
run.bold = True
run = p.add_run('Date: April.09.2026\n')
run.bold = True
run = p.add_run('Hours Worked: 9 Hours (8 hrs 10 min active, 50 min break)\n')
run.bold = True
run = p.add_run('Scheduled Hours Today: 11:00 pm (Apr 8) - 8:00 am (Apr 9) IST')
run.bold = True

# ── Remark ──
doc.add_heading('Remark', level=2)

p = doc.add_paragraph()
run = p.add_run('Algorithmic Trading System - Task Summary')
run.bold = True

doc.add_paragraph(
    'Wired Category B ML risk features (EVT, Monte Carlo, Hawkes, Temporal Transformer, sentiment) '
    'into the full ML pipeline (50 features), extended backtest engine with ML inference for impact '
    'measurement, fixed GPU server training errors, optimized backtest data loading with superset '
    'cache matching, and designed Bear/Risk Veto Agent architecture.'
)

# ── Key highlights ──
bullets = [
    'Expanded ML feature pipeline from 40 to 50 features across all models. Added 10 Category B '
    'risk/ML features: EVT VaR-99, EVT tail ratio, MC risk score, MC position scale, Hawkes intensity, '
    'TFT forecast bps, TFT confidence, sentiment mean, sentiment z-score, and regime encoding.',

    'Wired Category B direct entry_score modifiers into executor.py: EVT extreme tail (-2), EVT fat '
    'tail (-1), Monte Carlo high risk (-2/-1), Hawkes event clustering (-2/-1), and TFT direction '
    'conflict/confirmation (-1/+1). Added 4 new Category B ensemble voters to ML consensus voting.',

    'Added online autonomous learning hooks in executor main loop: EVT online_update() per bar for '
    'real-time tail risk tracking, Hawkes online_update() for event clustering adaptation, and TFT '
    'gradient descent per bar for forecast refinement.',

    'Fixed dead sentiment code in executor — replaced hardcoded placeholder values with real '
    'volume-weighted buying pressure proxy and rolling z-score computation per bar per asset.',

    'Extended training pipeline (train_all_models.py) to use local historical data (906K+ bars, '
    '8.6 years) instead of 20K bars from Binance API. Added load_local_data() function searching '
    'backtest_cache JSON, training_cache parquet, and data/ parquet. Added --use-local CLI flag.',

    'Fixed GPU server training crash: DatetimeArray casting error in parquet files. Created '
    '_parquet_to_data() helper with datetime64 detection and int64//10^6 millisecond conversion. '
    'Fixed single positional indexer out-of-bounds when bars=0 caused df.tail(0) empty DataFrame.',

    'Optimized backtest data_loader with superset cache matching — finds existing cached JSON files '
    'that cover the requested date range (e.g., 2016-2026 file covers 2017-2026 request). Added lazy '
    'exchange initialization to avoid Binance API connection when cache hit eliminates need for fetch.',

    'Built full ML inference pipeline into backtest engine (full_engine.py). Loads LightGBM, LSTM '
    'ensemble, EVT, Monte Carlo, Hawkes, and Temporal Transformer when --ml flag is set. Replicates '
    'executor.py scoring: individual model boosts/penalties, hard blocks, and 9-model ensemble voting '
    '(3+ SKIP = block, 2+ SKIP = score -3).',

    'Fixed LightGBM feature count mismatch in backtest ML inference — model trained on 50 features '
    'but code was truncating to 40 features ([:, :40] from old executor pattern). Now dynamically '
    'aligns with lgbm.num_feature().',

    'Trained all 14 ML models on server: BTC and ETH each with LightGBM (SKIP/TRADE binary), HMM '
    '(4-state regime), GARCH(1,1) volatility, Alpha Decay (optimal hold), RL Agent (Q-learning, '
    '300 states), LSTM Ensemble (LSTM+GRU+BiLSTM, 50 features), and PatchTST Transformer.',

    'Designed Bear/Risk Veto Agent plan — second LLM prompt to same Ollama model that argues AGAINST '
    'each trade. Scores risk 0-10 based on late entry, reversal patterns, volume divergence, key levels, '
    'and loss streaks. Risk >= 7 vetoes, 5-6 reduces position 50%.',

    'Diagnosed live trading halt: kill switch triggered by 69.84% testnet drawdown exceeding 10% limit. '
    'Confirmed halt is in-memory only — session PnL counters reset on restart. Provided restart commands '
    'to resume testnet trading.',
]

for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('Total Time Spent: 8 hr 10 min')
run.bold = True

# ── Task Table ──
tasks = [
    ['Task Type', 'Task Name', 'Completion Status', 'Hours worked on task today'],
    [
        'Feature\nEngineering',
        'Category B ML Feature Pipeline (50 features) - Expanded feature vector from 40 to 50 across '
        'LightGBM classifier, LSTM ensemble training, and executor inference. Added 10 Category B features: '
        'EVT VaR-99, tail ratio, MC risk score, MC position scale, Hawkes intensity, TFT forecast/confidence, '
        'sentiment mean/z-score, regime encoding. Updated extract_features(), compute_strategy_features(), '
        'and all n_features constants',
        'Completed', '1 hr 15 min'
    ],
    [
        'Feature\nIntegration',
        'Category B Direct Score Modifiers + Ensemble Voters (executor.py) - Wired EVT (-2/-1), '
        'Monte Carlo (-2/-1), Hawkes (-2/-1), TFT (-1/+1) as direct entry_score adjustments. Added '
        '4 Category B models to ML ensemble voting block (9 models total, 3+ SKIP = hard block)',
        'Completed', '45 min'
    ],
    [
        'Feature\nIntegration',
        'Online Autonomous Learning Hooks (executor.py) - Added per-bar online_update() calls for '
        'EVT tail risk, Hawkes event clustering, and Temporal Transformer gradient descent in main '
        'trading loop. Fixed dead sentiment code with real volume-weighted buying pressure proxy',
        'Completed', '35 min'
    ],
    [
        'Infrastructure\nImprovement',
        'Training Pipeline Local Data Support (train_all_models.py) - Added load_local_data() searching '
        '3 cache locations, _parquet_to_data() for datetime64 conversion, --use-local CLI flag. '
        'Enables training on 906K+ bars (8.6 years) instead of 20K API bars',
        'Completed', '40 min'
    ],
    [
        'Bug Fix &\nServer Support',
        'GPU Server Training Error Fixes - Fixed DatetimeArray cannot cast to float64 in parquet files '
        'with datetime64 detection + int64 conversion. Fixed single positional indexer out-of-bounds '
        'when bars=0 caused df.tail(0). Fixed Bybit fallback limit=0 edge case',
        'Completed', '30 min'
    ],
    [
        'Performance\nOptimization',
        'Backtest Data Loader Optimization (data_loader.py) - Added superset cache matching that finds '
        'cached files covering requested date range (eliminates redundant multi-hour Binance fetches). '
        'Added lazy exchange initialization (only connects if cache miss requires API fetch)',
        'Completed', '30 min'
    ],
    [
        'Feature\nDevelopment',
        'Backtest Engine ML Inference Pipeline (full_engine.py) - Built _run_ml_inference() method '
        'replicating executor.py ML pipeline: LightGBM binary gate, LSTM ensemble SKIP/TRADE, '
        'Category B score modifiers (EVT/MC/Hawkes/TFT), ML ensemble voting with hard blocks. '
        'Added --ml CLI flag to backtest.py. Tracks ML stats (blocks, boosts, penalties)',
        'Completed', '1 hr 10 min'
    ],
    [
        'Bug Fix',
        'LightGBM Feature Count Mismatch Fix (full_engine.py) - Model trained on 50 features but '
        'backtest code truncated to 40 ([:, :40] from old executor pattern). Fixed to dynamically '
        'align with lgbm.num_feature() using padding or truncation as needed',
        'Completed', '15 min'
    ],
    [
        'ML Model\nTraining',
        'Server Model Training (14/14) - Trained all models for BTC and ETH on server: LightGBM '
        '(BTC 67% acc, ETH 60% acc), HMM (4-state), GARCH(1,1), Alpha Decay, RL Agent (300+ states), '
        'LSTM Ensemble (3 sub-models, 50 features), PatchTST Transformer (55%/54.5% acc)',
        'Completed', '50 min'
    ],
    [
        'Architecture\n& Planning',
        'Bear/Risk Veto Agent Design - Designed dual-prompt architecture using same Ollama model: '
        'bull prompt (existing) + bear prompt (contrarian risk analyst). Risk score 0-10, veto at 7+, '
        'reduce at 5-6. Config-driven thresholds. Plan documented for implementation',
        'Completed', '25 min'
    ],
    [
        'Verification\n& Backtest',
        'Backtest Results Analysis (v15 Category B) - Ran 8.6-year backtests for BTC and ETH. '
        'BTC v15: 15,903 trades, WR=61.8%, PF=1.003, +$354. ETH v15: 17,141 trades, WR=61.0%, '
        'PF=1.018, +$2,842. Confirmed v15=v13/v14 without ML (proving ML wiring was needed)',
        'Completed', '30 min'
    ],
    [
        'System\nDiagnostics',
        'Live Trading Halt Investigation - Diagnosed kill switch triggered by 69.84% testnet '
        'drawdown (>10% limit). Traced halt mechanism through executor code. Confirmed in-memory '
        'state resets on restart. Provided restart procedure to resume trading',
        'Completed', '25 min'
    ],
]

table = doc.add_table(rows=len(tasks), cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
for row_idx, row_data in enumerate(tasks):
    row = table.rows[row_idx]
    for col_idx, cell_text in enumerate(row_data):
        cell = row.cells[col_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.font.size = Pt(9)
        if row_idx == 0:
            run.bold = True
            run.font.size = Pt(10)

# Save
output_path = r'C:\Users\convo\Downloads\Anudeep_Satya_Sai_EOD_April9.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
